from docx import Document
from docx.shared import Inches
import datetime
import os
def generate_insights_and_recommendations(topic_words_map, sentiment_score):
    """
    Dynamically generates business insights and actionable recommendations
    based on the dominant topic keywords and the sentiment score.
    Args:
        topic_words_map (dict): A dictionary where keys are topic IDs (int) 
                                and values are lists of top words (strings).
        sentiment_score (float): A value between 0.0 (Negative) and 1.0 (Positive).
    Returns:
        tuple: (List of insights strings, List of recommendations strings)
    """

    insights = []
    recommendations = []

    # Determine Sentiment Label
    if sentiment_score >= 0.6:
        sentiment_label = "Positive"
    elif sentiment_score <= 0.4:
        sentiment_label = "Negative"
    else:
        sentiment_label = "Neutral"

    # We define sets of keywords related to common themes in food reviews
    service_keywords = {'service', 'waiter', 'staff', 'manager', 'rude', 'slow', 'wait', 'friendly', 'attitude'}
    food_keywords = {'food', 'taste', 'flavor', 'delicious', 'bland', 'cold', 'menu', 'dish', 'meat', 'chicken', 'fresh'}
    ambience_keywords = {'place', 'atmosphere', 'music', 'noise', 'clean', 'dirty', 'table', 'seat', 'decor'}
    price_keywords = {'price', 'cost', 'expensive', 'cheap', 'value', 'bill', 'money', 'overpriced'}

    # Flatten all words from the provided map to find themes across the dominant topics
    all_topic_words = set()
    for words in topic_words_map.values():
        for w in words:
            all_topic_words.add(w.lower())

    # Check for overlapping themes using set intersection
    has_service = not service_keywords.isdisjoint(all_topic_words)
    has_food = not food_keywords.isdisjoint(all_topic_words)
    has_ambience = not ambience_keywords.isdisjoint(all_topic_words)
    has_price = not price_keywords.isdisjoint(all_topic_words)

    # Generate Insights & Recommendations based on Logic

    if sentiment_label == "Negative":
        insights.append("The overall sentiment is negative, indicating customer dissatisfaction.")
        
        if has_service:
            insights.append("Customers have raised concerns regarding staff behavior or service speed.")
            recommendations.append("Investigate staff training protocols and service speed. Address specific complaints about staff behavior immediately.")
        
        if has_food:
            insights.append("There are significant complaints regarding food quality, taste, or freshness.")
            recommendations.append("Review kitchen quality control and consistency. Check if specific dishes are repeatedly cited as poor.")
        
        if has_price:
            insights.append("Pricing or value for money is a major concern for customers.")
            recommendations.append("Re-evaluate portion sizes or menu pricing strategies to ensure customers feel they are getting good value.")

    elif sentiment_label == "Positive":
        insights.append("The customer feedback is largely positive, indicating a strong customer experience.")
        recommendations.append("Leverage these positive reviews in your marketing campaigns and social media.")
        
        if has_service:
            insights.append("Staff and service were highlighted positively.")
            recommendations.append("Identify and reward the staff members mentioned to boost team morale.")
        
        if has_food:
            insights.append("Customers specifically loved the food quality.")
            recommendations.append("Identify the most praised dishes and highlight them as 'Chef's Specials' or bestsellers.")

    else: # Neutral
        insights.append("The feedback is mixed or neutral, suggesting an 'okay' experience but room for improvement.")
        recommendations.append("Reach out to customers for more specific feedback to identify exactly what prevented a 5-star experience.")
        if has_ambience:
             insights.append("Ambience seems to be a talking point.")
             recommendations.append("Check if noise levels, lighting, or cleanliness are affecting the dining experience.")
    if not (has_service or has_food or has_ambience or has_price):
        insights.append("The review covers general topics without specific category keywords.")
        recommendations.append("Monitor more reviews to identify emerging trends that are not yet captured by standard categories.")

    return insights, recommendations

def build_docx_report(
    summary,
    sentiment_img_path,
    wordclouds,
    insights,
    recommendations,
    dominant_words=None,
    output_path = "assets/analysis_report.docx"
):
    """
    Build a DOCX report with summary, charts, wordclouds, insights, and recommendations.
    Saves the file to output_path and returns that path.
    """
    doc = Document()

    # Timestamp
    doc.add_paragraph(
        f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    # Summary
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(summary)

    # Visual Analysis
    doc.add_heading("Visual Analysis", level=2)

    if os.path.exists(sentiment_img_path):
        doc.add_paragraph("Sentiment Overview")
        doc.add_picture(sentiment_img_path, width=Inches(4))

    # Word Clouds
    if wordclouds:
        doc.add_heading("Word Clouds", level=2)
        for name, path in wordclouds.items():
            if os.path.exists(path):
                doc.add_paragraph(f"Word Cloud ({name})")
                doc.add_picture(path, width=Inches(4))

    # Dominant Topic Words
    if dominant_words:
        doc.add_heading("Dominant Topic Keywords", level=2)
        doc.add_paragraph(", ".join(dominant_words))

    # Insights
    doc.add_heading("Key Insights", level=2)
    if insights:
        for item in insights:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("No specific insights detected.")

    # Recommendations
    doc.add_heading("Actionable Recommendations", level=2)
    if recommendations:
        for item in recommendations:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("No specific recommendations generated.")

    # Save file
    doc.save(output_path)
    return output_path