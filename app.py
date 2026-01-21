# ======================================================
# STREAMLIT CONFIG
# ======================================================
import streamlit as st
st.set_page_config(
    page_title="AI Text Analysis Platform",
    page_icon="📊",
    layout="wide"
)

# ======================================================
# GLOBAL WEBSITE STYLES
# ======================================================
st.markdown("""
<style>
.stApp {
    background: linear-gradient(270deg,#020617,#0b1220,#020617);
    background-size: 500% 500%;
    animation: bgMove 20s ease infinite;
    color:#e5e7eb;
    font-family:Inter, sans-serif;
}
@keyframes bgMove {
    0% {background-position:0% 50%;}
    50% {background-position:100% 50%;}
    100% {background-position:0% 50%;}
}

.hero-title {
    font-size:3.4rem;
    font-weight:800;
    background: linear-gradient(90deg,#60a5fa,#3b82f6);
    -webkit-background-clip:text;
    -webkit-text-fill-color:transparent;
    text-align:center;
}

.hero-sub {
    text-align:center;
    max-width:760px;
    margin:auto;
    font-size:1.1rem;
    color:#cbd5f5;
}

.card {
    background:#0f172a;
    border:1px solid #1e293b;
    border-radius:16px;
    padding:1.5rem;
    min-height:160px;
    transition:0.3s;
}
.card:hover {
    transform:scale(1.03);
    border-color:#3b82f6;
}

.section-title {
    margin-top:2.5rem;
    font-size:1.4rem;
    font-weight:600;
}

.upload-box {
    background:#0f172a;
    border:2px dashed #334155;
    border-radius:16px;
    padding:2rem;
    text-align:center;
}
</style>
""", unsafe_allow_html=True)

# ======================================================
# IMPORTS
# ======================================================
import pandas as pd
import numpy as np
import joblib
from pathlib import Path
import plotly.express as px
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import time

# ======================================================
# PATHS & MODELS
# ======================================================
BASE_DIR = Path(__file__).resolve().parent
MODEL_DIR = BASE_DIR / "models"
TMP_DIR = BASE_DIR / "tmp"
TMP_DIR.mkdir(exist_ok=True)

@st.cache_resource
def load_models():
    return (
        joblib.load(MODEL_DIR / "tfidf_vectorizer.pkl"),
        joblib.load(MODEL_DIR / "nmf_model.pkl"),
        joblib.load(MODEL_DIR / "sentiment_model.pkl")
    )

tfidf, nmf_model, sentiment_model = load_models()

# ======================================================
# HELPERS
# ======================================================
def sentiment_label(score):
    if score >= 0.05: return "Positive"
    if score <= -0.05: return "Negative"
    return "Neutral"

def topic_table(df):
    words = tfidf.get_feature_names_out()
    rows, total = [], len(df)
    for i, comp in enumerate(nmf_model.components_):
        top = comp.argsort()[:-6:-1]
        kws = [words[j] for j in top]
        count = (df["Topic"] == i).sum()
        rows.append({
            "Topic ID": i,
            "Label": kws[0],
            "Top Keywords": ", ".join(kws),
            "Documents": count,
            "Share (%)": round(count/total*100, 2)
        })
    return pd.DataFrame(rows)

def executive_summary(df, topic_df, source):
    dom = topic_df.sort_values("Documents", ascending=False).iloc[0]
    return (
        f"The input source **{source}** contains **{len(df)} records**. "
        f"The dominant topic is **{dom['Label']}** contributing **{dom['Share (%)']}%** "
        f"of the data. Overall sentiment is **{df['Sentiment'].mode()[0]}**."
    )

def generate_docx(df, topic_df, sent_df, summary, images):
    doc = Document()
    doc.add_heading("AI Text Analysis Report", 1)
    doc.add_heading("Executive Summary", 2)
    doc.add_paragraph(summary)
    doc.add_heading("Topic Modeling Results", 2)
    doc.add_paragraph(topic_df.to_string(index=False))
    doc.add_heading("Sentiment Summary", 2)
    doc.add_paragraph(sent_df.to_string(index=False))
    doc.add_heading("Visual Insights", 2)
    for img in images:
        doc.add_picture(str(img), width=Inches(6))
    path = BASE_DIR / "AI_Text_Analysis_Report.docx"
    doc.save(path)
    return path

# ======================================================
# SESSION STATE
# ======================================================
st.session_state.setdefault("page", "home")
st.session_state.setdefault("df", None)

# ======================================================
# HOME PAGE
# ======================================================
if st.session_state.page == "home":

    st.markdown("<div class='hero-title'>AI Text Analysis Platform</div>", unsafe_allow_html=True)
    st.markdown("<div class='hero-sub'>AI-driven system for topic discovery, sentiment analysis and executive insights.</div>", unsafe_allow_html=True)

    st.markdown("### 👥 Who can use this platform?")
    cols = st.columns(4)
    users = [
        ("HR Teams","Employee feedback analysis"),
        ("Businesses","Customer reviews & complaints"),
        ("Researchers","Topic discovery & NLP"),
        ("Students","Learning AI & analytics")
    ]
    for c,(t,d) in zip(cols,users):
        with c:
            st.markdown(f"<div class='card'><b>{t}</b><br/>{d}</div>", unsafe_allow_html=True)

    st.markdown("### 📘 How to Use")
    st.markdown("""
    1. Upload CSV or enter text  
    2. Run AI analysis  
    3. Explore insights  
    4. Download report  
    """)

    if st.button("🚀 Get Started"):
        st.session_state.page = "analysis"
        st.rerun()

# ======================================================
# ANALYSIS PAGE
# ======================================================
if st.session_state.page == "analysis":

    if st.button("⬅ Back to Home"):
        st.session_state.page = "home"
        st.rerun()

    st.markdown("## 📊 Analysis Dashboard")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("Upload CSV</div>", unsafe_allow_html=True)
        csv_file = st.file_uploader("", type=["csv"])
    with c2:
        st.markdown("Enter Text</div>", unsafe_allow_html=True)
        text_input = st.text_area("", height=150)

    if st.button("▶ Analyze"):

        if csv_file:
            df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
            df["combined_text"] = df.select_dtypes("object").astype(str).agg(" ".join, axis=1)
            source = csv_file.name

        elif text_input.strip():
            df = pd.DataFrame({"combined_text":[text_input]})
            source = "Text Input"

        else:
            st.warning("Please upload CSV or enter text")
            st.stop()

        vecs = tfidf.transform(df["combined_text"])
        df["Topic"] = np.argmax(nmf_model.transform(vecs), axis=1)
        df["Sentiment"] = df["combined_text"].apply(
            lambda x: sentiment_label(sentiment_model.polarity_scores(x)["compound"])
        )

        st.session_state.df = df
        st.session_state.source = source

    # ======================================================
    # RESULTS
    # ======================================================
    if st.session_state.df is not None:
        df = st.session_state.df
        topic_df = topic_table(df)
        sent_df = df["Sentiment"].value_counts().reset_index()
        sent_df.columns = ["Sentiment","Count"]

        st.markdown("### 🧠 Executive Summary")
        summary = executive_summary(df, topic_df, st.session_state.source)
        st.info(summary)

        st.markdown("### 📌 Topic Modeling Table")
        st.dataframe(topic_df, use_container_width=True)

        colA, colB = st.columns(2)
        with colA:
            st.markdown("**Topic Share**")
            st.plotly_chart(px.pie(topic_df, names="Label", values="Share (%)", hole=0.55))
        with colB:
            st.markdown("**Documents per Topic**")
            st.plotly_chart(px.bar(topic_df, x="Topic ID", y="Documents", labels={"Documents":"Count"}))

        colC, colD = st.columns(2)
        with colC:
            st.markdown("**Sentiment Distribution**")
            st.plotly_chart(px.pie(sent_df, names="Sentiment", values="Count", hole=0.6))
        with colD:
            st.markdown("**Sentiment Ratio**")
            st.plotly_chart(px.bar(sent_df, x="Sentiment", y="Count"))

        st.markdown("### ☁ Word Cloud")
        wc_path = TMP_DIR / "wordcloud.png"
        WordCloud(width=900, height=350, background_color="#020617") \
            .generate(" ".join(df["combined_text"])).to_file(wc_path)
        st.image(wc_path)

        st.markdown("### 📊 Topic × Sentiment")
        st.plotly_chart(
            px.bar(
                df.groupby(["Topic","Sentiment"]).size().reset_index(name="Count"),
                x="Topic", y="Count", color="Sentiment", barmode="stack"
            )
        )

        st.markdown("### 🔍 Auto Insights")
        st.success(
            f"Dominant topic: {topic_df.iloc[0]['Label']} | "
            f"Overall sentiment: {df['Sentiment'].mode()[0]}"
        )

        st.divider()

        # CREATE COLUMNS FIRST (FIX)
        colD1, colD2 = st.columns([2, 1])

        # ==========================
        # DOWNLOAD FULL REPORT
        # ==========================
        with colD1:
            if st.button("⬇ Download Full Report", type="primary"):

                imgs = []

                def save_fig(fig, name):
                    path = TMP_DIR / name
                    fig.savefig(path, bbox_inches="tight")
                    plt.close(fig)
                    imgs.append(path)

                # Topic Share Pie
                fig, ax = plt.subplots(figsize=(6,4))
                ax.pie(
                    topic_df["Share (%)"],
                    labels=topic_df["Label"],
                    autopct="%1.1f%%",
                    startangle=140
                )
                ax.set_title("Topic Share Distribution")
                save_fig(fig, "topic_share.png")

                # Documents per Topic
                fig, ax = plt.subplots(figsize=(6,4))
                ax.bar(topic_df["Topic ID"], topic_df["Documents"])
                ax.set_title("Documents per Topic")
                save_fig(fig, "documents_per_topic.png")

                # Sentiment Distribution
                fig, ax = plt.subplots(figsize=(6,4))
                ax.pie(
                    sent_df["Count"],
                    labels=sent_df["Sentiment"],
                    autopct="%1.1f%%"
                )
                ax.set_title("Sentiment Distribution")
                save_fig(fig, "sentiment_distribution.png")

                # Topic × Sentiment
                pivot = df.groupby(["Topic","Sentiment"]).size().unstack(fill_value=0)
                fig, ax = plt.subplots(figsize=(7,4))
                pivot.plot(kind="bar", stacked=True, ax=ax)
                ax.set_title("Topic × Sentiment Analysis")
                save_fig(fig, "topic_sentiment.png")

                # Word Cloud
                wc_path = TMP_DIR / "wordcloud.png"
                WordCloud(
                    width=900,
                    height=350,
                    background_color="#020617"
                ).generate(" ".join(df["combined_text"])).to_file(wc_path)
                imgs.append(wc_path)

                # Generate DOCX
                path = generate_docx(
                    df,
                    topic_df,
                    sent_df,
                    summary,
                    imgs
                )

                with open(path, "rb") as f:
                    st.download_button(
                        "📄 Download Report (DOCX)",
                        f,
                        file_name="AI_Text_Analysis_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        # ==========================
        # RESET BUTTON
        # ==========================
        with colD2:
            if st.button("🔄 Reset"):
                st.session_state.df = None
                st.session_state.page = "home"
                st.rerun()
